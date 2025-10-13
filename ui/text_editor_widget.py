# --- NEW FEATURE: Text Editor ---
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
                             QTextEdit, QFileDialog, QMessageBox, QToolBar,
                             QFontComboBox, QSpinBox, QDialog, QLabel,
                             QDialogButtonBox, QMenu, QDoubleSpinBox, QComboBox,
                             QColorDialog)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import (QFont, QTextCharFormat, QTextCursor, QPageSize, QPageLayout, 
                         QTextTableFormat, QAction, QColor, QTextListFormat)
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import os

class TextEditorWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.current_file = None
        self.setup_ui()
    
    def setup_ui(self):
        layout = QVBoxLayout()
        
        # Toolbar
        toolbar_layout = QHBoxLayout()
        
        # Font selection
        self.font_combo = QFontComboBox()
        self.font_combo.currentFontChanged.connect(self.change_font)
        toolbar_layout.addWidget(self.font_combo)
        
        # Font size
        self.font_size = QSpinBox()
        self.font_size.setValue(12)
        self.font_size.setMinimum(6)
        self.font_size.setMaximum(72)
        self.font_size.valueChanged.connect(self.change_font_size)
        toolbar_layout.addWidget(self.font_size)
        
        # Bold button
        bold_btn = QPushButton('B')
        bold_btn.setFont(QFont('Arial', 12, QFont.Weight.Bold))
        bold_btn.setCheckable(True)
        bold_btn.clicked.connect(self.toggle_bold)
        bold_btn.setFixedWidth(40)
        toolbar_layout.addWidget(bold_btn)
        
        # Italic button
        italic_btn = QPushButton('I')
        italic_btn.setFont(QFont('Arial', 12, QFont.Weight.Normal))
        italic_btn.font().setItalic(True)
        italic_btn.setCheckable(True)
        italic_btn.clicked.connect(self.toggle_italic)
        italic_btn.setFixedWidth(40)
        toolbar_layout.addWidget(italic_btn)
        
        # Underline button
        underline_btn = QPushButton('U')
        underline_btn.setFont(QFont('Arial', 12, QFont.Weight.Normal))
        underline_btn.font().setUnderline(True)
        underline_btn.setCheckable(True)
        underline_btn.clicked.connect(self.toggle_underline)
        underline_btn.setFixedWidth(40)
        toolbar_layout.addWidget(underline_btn)
        
        toolbar_layout.addSpacing(20)
        
        # Alignment buttons
        align_left_btn = QPushButton('⬅')
        align_left_btn.clicked.connect(self.align_left)
        align_left_btn.setFixedWidth(40)
        toolbar_layout.addWidget(align_left_btn)
        
        align_center_btn = QPushButton('↔')
        align_center_btn.clicked.connect(self.align_center)
        align_center_btn.setFixedWidth(40)
        toolbar_layout.addWidget(align_center_btn)
        
        align_right_btn = QPushButton('➡')
        align_right_btn.clicked.connect(self.align_right)
        align_right_btn.setFixedWidth(40)
        toolbar_layout.addWidget(align_right_btn)
        
        toolbar_layout.addSpacing(20)
        
        text_color_btn = QPushButton('🎨 لون النص')
        text_color_btn.clicked.connect(self.change_text_color)
        toolbar_layout.addWidget(text_color_btn)
        
        bg_color_btn = QPushButton('🖌️ لون الخلفية')
        bg_color_btn.clicked.connect(self.change_background_color)
        toolbar_layout.addWidget(bg_color_btn)
        
        toolbar_layout.addSpacing(20)
        
        bullet_btn = QPushButton('• قائمة منقطة')
        bullet_btn.clicked.connect(self.insert_bullet_list)
        toolbar_layout.addWidget(bullet_btn)
        
        numbered_btn = QPushButton('1. قائمة مرقمة')
        numbered_btn.clicked.connect(self.insert_numbered_list)
        toolbar_layout.addWidget(numbered_btn)
        
        toolbar_layout.addStretch()
        layout.addLayout(toolbar_layout)
        
        # Text editor
        self.text_edit = QTextEdit()
        self.text_edit.setFont(QFont('Arial', 12))
        self.text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.show_context_menu)
        layout.addWidget(self.text_edit)
        
        # Action buttons
        btn_layout = QHBoxLayout()
        
        new_btn = QPushButton('📄 جديد')
        new_btn.clicked.connect(self.new_document)
        btn_layout.addWidget(new_btn)
        
        open_btn = QPushButton('📂 فتح')
        open_btn.clicked.connect(self.open_file)
        btn_layout.addWidget(open_btn)
        
        save_txt_btn = QPushButton('💾 حفظ TXT')
        save_txt_btn.clicked.connect(self.save_as_txt)
        btn_layout.addWidget(save_txt_btn)
        
        save_docx_btn = QPushButton('💾 حفظ DOCX')
        save_docx_btn.clicked.connect(self.save_as_docx)
        btn_layout.addWidget(save_docx_btn)
        
        print_btn = QPushButton('🖨️ طباعة')
        print_btn.clicked.connect(self.print_document)
        btn_layout.addWidget(print_btn)
        
        # --- NEW (زر إدراج جدول) ---
        insert_table_btn = QPushButton('📊 إدراج جدول')
        insert_table_btn.clicked.connect(self.insert_table)
        btn_layout.addWidget(insert_table_btn)
        
        edit_table_btn = QPushButton('✏️ تحرير جدول')
        edit_table_btn.clicked.connect(self.edit_current_table)
        btn_layout.addWidget(edit_table_btn)
        
        undo_btn = QPushButton('↶ تراجع')
        undo_btn.clicked.connect(self.text_edit.undo)
        btn_layout.addWidget(undo_btn)
        
        redo_btn = QPushButton('↷ إعادة')
        redo_btn.clicked.connect(self.text_edit.redo)
        btn_layout.addWidget(redo_btn)
        
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        self.setLayout(layout)
    
    def change_font(self, font):
        self.text_edit.setCurrentFont(font)
    
    def change_font_size(self, size):
        self.text_edit.setFontPointSize(size)
    
    def toggle_bold(self):
        fmt = QTextCharFormat()
        if self.text_edit.fontWeight() == QFont.Weight.Bold:
            fmt.setFontWeight(QFont.Weight.Normal)
        else:
            fmt.setFontWeight(QFont.Weight.Bold)
        self.text_edit.mergeCurrentCharFormat(fmt)
    
    def toggle_italic(self):
        fmt = QTextCharFormat()
        fmt.setFontItalic(not self.text_edit.fontItalic())
        self.text_edit.mergeCurrentCharFormat(fmt)
    
    def toggle_underline(self):
        fmt = QTextCharFormat()
        fmt.setFontUnderline(not self.text_edit.fontUnderline())
        self.text_edit.mergeCurrentCharFormat(fmt)
    
    def align_left(self):
        self.text_edit.setAlignment(Qt.AlignmentFlag.AlignLeft)
    
    def align_center(self):
        self.text_edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
    
    def align_right(self):
        self.text_edit.setAlignment(Qt.AlignmentFlag.AlignRight)
    
    def change_text_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            self.text_edit.mergeCurrentCharFormat(fmt)
    
    def change_background_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            self.text_edit.mergeCurrentCharFormat(fmt)
    
    def insert_bullet_list(self):
        cursor = self.text_edit.textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.Style.ListDisc)
        cursor.createList(list_format)
    
    def insert_numbered_list(self):
        cursor = self.text_edit.textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.Style.ListDecimal)
        cursor.createList(list_format)
    
    def new_document(self):
        if self.text_edit.document().isModified():
            reply = QMessageBox.question(
                self,
                'حفظ التغييرات؟',
                'هل تريد حفظ التغييرات قبل إنشاء مستند جديد؟',
                QMessageBox.StandardButton.Yes | 
                QMessageBox.StandardButton.No | 
                QMessageBox.StandardButton.Cancel
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.save_as_txt()
            elif reply == QMessageBox.StandardButton.Cancel:
                return
        
        self.text_edit.clear()
        self.current_file = None
    
    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            'فتح ملف',
            '',
            'Text Files (*.txt);;Word Documents (*.docx);;All Files (*)'
        )
        
        if file_path:
            try:
                if file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        self.text_edit.setPlainText(f.read())
                elif file_path.endswith('.docx'):
                    doc = Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    self.text_edit.setPlainText('\n'.join(full_text))
                
                self.current_file = file_path
                QMessageBox.information(self, 'نجح', 'تم فتح الملف بنجاح')
            except Exception as e:
                QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء فتح الملف:\n{str(e)}')
    
    def save_as_txt(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            'حفظ كملف نصي',
            self.current_file if self.current_file and self.current_file.endswith('.txt') else 'document.txt',
            'Text Files (*.txt)'
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.text_edit.toPlainText())
                
                self.current_file = file_path
                self.text_edit.document().setModified(False)
                QMessageBox.information(self, 'نجح', 'تم حفظ الملف بنجاح')
            except Exception as e:
                QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء الحفظ:\n{str(e)}')
    
    def save_as_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            'حفظ كمستند Word',
            self.current_file if self.current_file and self.current_file.endswith('.docx') else 'document.docx',
            'Word Documents (*.docx)'
        )
        
        if file_path:
            try:
                doc = Document()
                
                # --- FIX (تحويل الجداول من QTextEdit إلى Word) ---
                # Get HTML content to preserve tables
                html_content = self.text_edit.toHtml()
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Process HTML elements
                for element in soup.find_all(['p', 'table']):
                    if element.name == 'table':
                        # Extract table data
                        rows_data = []
                        for tr in element.find_all('tr'):
                            row_data = []
                            for td in tr.find_all('td'):
                                row_data.append(td.get_text(strip=True))
                            if row_data:
                                rows_data.append(row_data)
                        
                        # Create Word table
                        if rows_data:
                            num_rows = len(rows_data)
                            num_cols = len(rows_data[0]) if rows_data else 0
                            
                            if num_cols > 0:
                                table = doc.add_table(rows=num_rows, cols=num_cols)
                                table.style = 'Table Grid'
                                
                                # Fill table with data
                                for i, row_data in enumerate(rows_data):
                                    for j, cell_text in enumerate(row_data):
                                        if j < len(table.rows[i].cells):
                                            table.rows[i].cells[j].text = cell_text
                                            # Set RTL for Arabic text
                                            for paragraph in table.rows[i].cells[j].paragraphs:
                                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    elif element.name == 'p':
                        # Add paragraph
                        text = element.get_text(strip=True)
                        if text:
                            para = doc.add_paragraph(text)
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            doc.add_paragraph()
                
                doc.save(file_path)
                self.current_file = file_path
                self.text_edit.document().setModified(False)
                QMessageBox.information(self, 'نجح', 'تم حفظ الملف بنجاح مع الجداول')
            except Exception as e:
                QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء الحفظ:\n{str(e)}')
    
    def insert_table(self):
        # --- NEW (إدراج جدول في محرر النصوص) ---
        dialog = QDialog(self)
        dialog.setWindowTitle('إدراج جدول')
        dialog.setFixedWidth(350)
        
        layout = QVBoxLayout()
        
        rows_layout = QHBoxLayout()
        rows_layout.addWidget(QLabel('عدد الصفوف:'))
        rows_input = QSpinBox()
        rows_input.setMinimum(1)
        rows_input.setMaximum(50)
        rows_input.setValue(3)
        rows_layout.addWidget(rows_input)
        layout.addLayout(rows_layout)
        
        cols_layout = QHBoxLayout()
        cols_layout.addWidget(QLabel('عدد الأعمدة:'))
        cols_input = QSpinBox()
        cols_input.setMinimum(1)
        cols_input.setMaximum(20)
        cols_input.setValue(3)
        cols_layout.addWidget(cols_input)
        layout.addLayout(cols_layout)
        
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            rows = rows_input.value()
            cols = cols_input.value()
            
            # Insert table using QTextTable
            cursor = self.text_edit.textCursor()
            table_format = QTextTableFormat()
            table_format.setBorderStyle(QTextTableFormat.BorderStyle.BorderStyle_Solid)
            table_format.setCellPadding(5)
            table_format.setCellSpacing(0)
            table = cursor.insertTable(rows, cols, table_format)
            
            QMessageBox.information(self, 'نجح', f'تم إدراج جدول {rows}×{cols} بنجاح')
    
    def print_document(self):
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            dialog = QPrintDialog(printer, self)
            
            if dialog.exec() == QPrintDialog.DialogCode.Accepted:
                self.text_edit.document().print(printer)
                QMessageBox.information(self, 'نجح', 'تم الطباعة بنجاح')
        except Exception as e:
            QMessageBox.critical(self, 'خطأ', f'حدث خطأ أثناء الطباعة:\n{str(e)}')
    
    def show_context_menu(self, position):
        cursor = self.text_edit.cursorForPosition(position)
        table = cursor.currentTable()
        
        menu = QMenu(self)
        
        if table:
            add_row_action = QAction('إضافة صف', self)
            add_row_action.triggered.connect(lambda: self.add_table_row(table, cursor))
            menu.addAction(add_row_action)
            
            remove_row_action = QAction('حذف صف', self)
            remove_row_action.triggered.connect(lambda: self.remove_table_row(table, cursor))
            menu.addAction(remove_row_action)
            
            menu.addSeparator()
            
            add_col_action = QAction('إضافة عمود', self)
            add_col_action.triggered.connect(lambda: self.add_table_column(table, cursor))
            menu.addAction(add_col_action)
            
            remove_col_action = QAction('حذف عمود', self)
            remove_col_action.triggered.connect(lambda: self.remove_table_column(table, cursor))
            menu.addAction(remove_col_action)
            
            menu.addSeparator()
            
            edit_table_action = QAction('تحرير خصائص الجدول', self)
            edit_table_action.triggered.connect(lambda: self.edit_table_properties(table))
            menu.addAction(edit_table_action)
            
            menu.addSeparator()
            
            merge_cells_action = QAction('دمج الخلايا المحددة', self)
            merge_cells_action.triggered.connect(lambda: self.merge_selected_cells(table, cursor))
            menu.addAction(merge_cells_action)
            
            split_cell_action = QAction('فك دمج الخلية', self)
            split_cell_action.triggered.connect(lambda: self.split_current_cell(table, cursor))
            menu.addAction(split_cell_action)
        else:
            default_action = QAction('لا يوجد جدول في موضع المؤشر', self)
            default_action.setEnabled(False)
            menu.addAction(default_action)
        
        menu.exec(self.text_edit.mapToGlobal(position))
    
    def add_table_row(self, table, cursor):
        cell = table.cellAt(cursor)
        row = cell.row()
        table.insertRows(row + 1, 1)
        QMessageBox.information(self, 'نجح', 'تم إضافة صف جديد')
    
    def remove_table_row(self, table, cursor):
        if table.rows() <= 1:
            QMessageBox.warning(self, 'تحذير', 'لا يمكن حذف الصف الوحيد في الجدول')
            return
        
        cell = table.cellAt(cursor)
        row = cell.row()
        table.removeRows(row, 1)
        QMessageBox.information(self, 'نجح', 'تم حذف الصف')
    
    def add_table_column(self, table, cursor):
        cell = table.cellAt(cursor)
        col = cell.column()
        table.insertColumns(col + 1, 1)
        QMessageBox.information(self, 'نجح', 'تم إضافة عمود جديد')
    
    def remove_table_column(self, table, cursor):
        if table.columns() <= 1:
            QMessageBox.warning(self, 'تحذير', 'لا يمكن حذف العمود الوحيد في الجدول')
            return
        
        cell = table.cellAt(cursor)
        col = cell.column()
        table.removeColumns(col, 1)
        QMessageBox.information(self, 'نجح', 'تم حذف العمود')
    
    def edit_table_properties(self, table):
        dialog = QDialog(self)
        dialog.setWindowTitle('تحرير خصائص الجدول')
        dialog.setFixedWidth(400)
        
        layout = QVBoxLayout()
        
        border_layout = QHBoxLayout()
        border_layout.addWidget(QLabel('نمط الإطار:'))
        border_style_combo = QComboBox()
        border_style_combo.addItems(['بدون إطار', 'خط مفرد', 'خط مزدوج', 'خط منقط', 'خط متقطع'])
        border_layout.addWidget(border_style_combo)
        layout.addLayout(border_layout)
        
        current_format = table.format()
        current_border = current_format.borderStyle()
        
        if current_border == QTextTableFormat.BorderStyle.BorderStyle_None:
            border_style_combo.setCurrentIndex(0)
        elif current_border == QTextTableFormat.BorderStyle.BorderStyle_Solid:
            border_style_combo.setCurrentIndex(1)
        elif current_border == QTextTableFormat.BorderStyle.BorderStyle_Double:
            border_style_combo.setCurrentIndex(2)
        elif current_border == QTextTableFormat.BorderStyle.BorderStyle_Dotted:
            border_style_combo.setCurrentIndex(3)
        elif current_border == QTextTableFormat.BorderStyle.BorderStyle_Dashed:
            border_style_combo.setCurrentIndex(4)
        
        width_layout = QHBoxLayout()
        width_layout.addWidget(QLabel('عرض الإطار:'))
        border_width_spin = QDoubleSpinBox()
        border_width_spin.setMinimum(0)
        border_width_spin.setMaximum(10)
        border_width_spin.setValue(current_format.border())
        border_width_spin.setSingleStep(0.5)
        width_layout.addWidget(border_width_spin)
        layout.addLayout(width_layout)
        
        padding_layout = QHBoxLayout()
        padding_layout.addWidget(QLabel('المسافة الداخلية للخلايا:'))
        cell_padding_spin = QDoubleSpinBox()
        cell_padding_spin.setMinimum(0)
        cell_padding_spin.setMaximum(20)
        cell_padding_spin.setValue(current_format.cellPadding())
        cell_padding_spin.setSingleStep(1)
        padding_layout.addWidget(cell_padding_spin)
        layout.addLayout(padding_layout)
        
        spacing_layout = QHBoxLayout()
        spacing_layout.addWidget(QLabel('المسافة بين الخلايا:'))
        cell_spacing_spin = QDoubleSpinBox()
        cell_spacing_spin.setMinimum(0)
        cell_spacing_spin.setMaximum(20)
        cell_spacing_spin.setValue(current_format.cellSpacing())
        cell_spacing_spin.setSingleStep(1)
        spacing_layout.addWidget(cell_spacing_spin)
        layout.addLayout(spacing_layout)
        
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_format = QTextTableFormat()
            
            border_styles = [
                QTextTableFormat.BorderStyle.BorderStyle_None,
                QTextTableFormat.BorderStyle.BorderStyle_Solid,
                QTextTableFormat.BorderStyle.BorderStyle_Double,
                QTextTableFormat.BorderStyle.BorderStyle_Dotted,
                QTextTableFormat.BorderStyle.BorderStyle_Dashed
            ]
            new_format.setBorderStyle(border_styles[border_style_combo.currentIndex()])
            new_format.setBorder(border_width_spin.value())
            new_format.setCellPadding(cell_padding_spin.value())
            new_format.setCellSpacing(cell_spacing_spin.value())
            
            table.setFormat(new_format)
            QMessageBox.information(self, 'نجح', 'تم تحديث خصائص الجدول بنجاح')
    
    def edit_current_table(self):
        cursor = self.text_edit.textCursor()
        table = cursor.currentTable()
        
        if table:
            self.edit_table_properties(table)
        else:
            QMessageBox.warning(self, 'تحذير', 'الرجاء وضع المؤشر داخل الجدول المراد تحريره')
    
    def merge_selected_cells(self, table, cursor):
        cell = table.cellAt(cursor)
        if not cell.isValid():
            QMessageBox.warning(self, 'خطأ', 'لا يمكن تحديد الخلية')
            return
        
        dialog = QDialog(self)
        dialog.setWindowTitle('دمج الخلايا')
        dialog.setFixedWidth(350)
        
        layout = QVBoxLayout()
        
        info_label = QLabel(f'الخلية الحالية: صف {cell.row() + 1}, عمود {cell.column() + 1}')
        layout.addWidget(info_label)
        
        rows_layout = QHBoxLayout()
        rows_layout.addWidget(QLabel('عدد الصفوف للدمج:'))
        rows_spin = QSpinBox()
        rows_spin.setMinimum(1)
        rows_spin.setMaximum(table.rows() - cell.row())
        rows_spin.setValue(1)
        rows_layout.addWidget(rows_spin)
        layout.addLayout(rows_layout)
        
        cols_layout = QHBoxLayout()
        cols_layout.addWidget(QLabel('عدد الأعمدة للدمج:'))
        cols_spin = QSpinBox()
        cols_spin.setMinimum(1)
        cols_spin.setMaximum(table.columns() - cell.column())
        cols_spin.setValue(1)
        cols_layout.addWidget(cols_spin)
        layout.addLayout(cols_layout)
        
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            num_rows = rows_spin.value()
            num_cols = cols_spin.value()
            
            if num_rows > 1 or num_cols > 1:
                table.mergeCells(cell.row(), cell.column(), num_rows, num_cols)
                QMessageBox.information(self, 'نجح', f'تم دمج {num_rows}×{num_cols} خلايا بنجاح')
            else:
                QMessageBox.information(self, 'تنبيه', 'لم يتم تحديد خلايا للدمج')
    
    def split_current_cell(self, table, cursor):
        cell = table.cellAt(cursor)
        if not cell.isValid():
            QMessageBox.warning(self, 'خطأ', 'لا يمكن تحديد الخلية')
            return
        
        row_span = cell.rowSpan()
        col_span = cell.columnSpan()
        
        if row_span > 1 or col_span > 1:
            dialog = QDialog(self)
            dialog.setWindowTitle('فك دمج الخلية')
            dialog.setFixedWidth(350)
            
            layout = QVBoxLayout()
            
            info_label = QLabel(f'''
            <div style="text-align: right; direction: rtl;">
            <p>الخلية الحالية مدمجة:</p>
            <p><b>عدد الصفوف:</b> {row_span}</p>
            <p><b>عدد الأعمدة:</b> {col_span}</p>
            <p>سيتم فك دمج هذه الخلية إلى {row_span * col_span} خلية منفصلة</p>
            </div>
            ''')
            info_label.setTextFormat(Qt.TextFormat.RichText)
            layout.addWidget(info_label)
            
            buttons = QDialogButtonBox(
                QDialogButtonBox.StandardButton.Ok | 
                QDialogButtonBox.StandardButton.Cancel
            )
            buttons.accepted.connect(dialog.accept)
            buttons.rejected.connect(dialog.reject)
            layout.addWidget(buttons)
            
            dialog.setLayout(layout)
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                table.splitCell(cell.row(), cell.column(), 1, 1)
                QMessageBox.information(self, 'نجح', 'تم فك دمج الخلية بنجاح')
        else:
            QMessageBox.information(self, 'تنبيه', 'هذه الخلية غير مدمجة')
